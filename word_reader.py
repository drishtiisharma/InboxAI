from docx import Document

def extract_text_from_docx(file_path):
<<<<<<< HEAD
    doc = Document(file_path)
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    return "\n".join(text)
=======
    try:
        doc = Document(file_path)
        text = []

        # Extract normal paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text.append(para.text.strip())

        # Extract tables (important for reports/invoices)
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if row_text:
                    text.append(" | ".join(row_text))

        return "\n".join(text)

    except Exception as e:
        return f"[ERROR reading document: {str(e)}]"
>>>>>>> backend
